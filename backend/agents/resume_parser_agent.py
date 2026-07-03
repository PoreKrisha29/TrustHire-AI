"""
agents/resume_parser_agent.py

Parse resume files (PDF or DOCX) into a structured dict using:
  - PyMuPDF (fitz) for PDF text extraction
  - python-docx for DOCX text extraction
  - Google Gemini 1.5 Flash for structured information extraction

Output schema:
    {
        "name": str,
        "email": str,
        "phone": str,
        "summary": str,
        "skills": [str, ...],           # flat list of skill strings
        "experience": [
            {
                "company": str,
                "role": str,
                "duration": str,
                "bullets": [str, ...]
            },
            ...
        ],
        "education": [
            {
                "institution": str,
                "degree": str,
                "year": str
            },
            ...
        ],
        "projects": [
            {
                "title": str,
                "description": str,
                "tech": [str, ...]
            },
            ...
        ]
    }
"""

import io
import json
import re

import google.generativeai as genai
from django.conf import settings


# ---------------------------------------------------------------------------
# Gemini client (lazy-initialised)
# ---------------------------------------------------------------------------

_gemini_model = None


def _get_model():
    global _gemini_model
    if _gemini_model is None:
        genai.configure(api_key=settings.GEMINI_API_KEY)
        _gemini_model = genai.GenerativeModel(settings.GEMINI_MODEL)
    return _gemini_model


# ---------------------------------------------------------------------------
# Text extractors
# ---------------------------------------------------------------------------

def _extract_text_pdf(file_bytes: bytes) -> str:
    """Extract all text from a PDF using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = [page.get_text("text") for page in doc]
        doc.close()
        return "\n".join(pages)
    except Exception as exc:
        raise ValueError(f"PDF text extraction failed: {exc}") from exc


def _extract_text_docx(file_bytes: bytes) -> str:
    """Extract all text from a DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as exc:
        raise ValueError(f"DOCX text extraction failed: {exc}") from exc


# ---------------------------------------------------------------------------
# Gemini extraction
# ---------------------------------------------------------------------------

_EXTRACTION_PROMPT = """
You are a professional resume parser. Extract structured information from the resume text below.

Return ONLY a valid JSON object — no markdown fences, no explanation.

Schema to follow exactly:
{{
  "name": "string",
  "email": "string",
  "phone": "string",
  "summary": "string (professional summary or objective, 2-4 sentences)",
  "skills": ["skill1", "skill2"],
  "experience": [
    {{
      "company": "string",
      "role": "string",
      "duration": "string (e.g. Jan 2022 – Present)",
      "bullets": ["bullet1", "bullet2"]
    }}
  ],
  "education": [
    {{
      "institution": "string",
      "degree": "string",
      "year": "string"
    }}
  ],
  "projects": [
    {{
      "title": "string",
      "description": "string",
      "tech": ["tech1", "tech2"]
    }}
  ]
}}

Rules:
- skills must be a flat list of plain strings (e.g. ["Python", "React", "PostgreSQL"])
- If a field is not found, use an empty string or empty list
- Do NOT add any extra keys
- Do NOT wrap in markdown code blocks

Resume text:
---
{resume_text}
---
"""


def _parse_with_gemini(raw_text: str) -> dict:
    """Send raw resume text to Gemini and return parsed structured dict."""
    model = _get_model()
    prompt = _EXTRACTION_PROMPT.format(resume_text=raw_text[:12000])  # token safety

    response = model.generate_content(
        prompt,
        generation_config=genai.GenerationConfig(
            temperature=0.1,
            response_mime_type="application/json",
        ),
    )

    text = response.text.strip()
    # Strip accidental markdown fences
    text = re.sub(r"^```(?:json)?\s*", "", text)
    text = re.sub(r"\s*```$", "", text)

    try:
        data = json.loads(text)
    except json.JSONDecodeError as exc:
        raise ValueError(f"Gemini returned invalid JSON: {exc}\nRaw: {text[:500]}") from exc

    # Normalise: ensure skills is always list[str]
    skills = data.get("skills", [])
    if skills and isinstance(skills[0], dict):
        data["skills"] = [s.get("skill", str(s)) for s in skills]

    return data


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def parse_pdf(file_bytes: bytes) -> dict:
    """
    Parse a PDF resume file into a structured dict.

    Args:
        file_bytes: Raw bytes of the PDF file.

    Returns:
        Structured resume dict matching the schema above.
    """
    raw_text = _extract_text_pdf(file_bytes)
    if not raw_text.strip():
        raise ValueError("No readable text found in the PDF. It may be image-based.")
    return _parse_with_gemini(raw_text)


def parse_docx(file_bytes: bytes) -> dict:
    """
    Parse a DOCX resume file into a structured dict.

    Args:
        file_bytes: Raw bytes of the DOCX file.

    Returns:
        Structured resume dict matching the schema above.
    """
    raw_text = _extract_text_docx(file_bytes)
    if not raw_text.strip():
        raise ValueError("No readable text found in the DOCX.")
    return _parse_with_gemini(raw_text)
